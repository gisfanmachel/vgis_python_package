"""
#!/usr/bin/python3.9
# -*- coding: utf-8 -*-
@Project :pythonCodeSnippet
@File    :wordTools.py
@IDE     :PyCharm
@Author  :chenxw
@Date    :2023/9/22 14:51
@Descr:
"""
import os

from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn
from docx.shared import Cm
from docx.shared import Pt
import win32com.client
from pathlib import Path
import pythoncom

class WordHelper:

    @staticmethod
    # 需要安装wps
    def wps2docx(wps_file_path, save_path):
        pythoncom.CoInitialize()
        # 使用 WPS 打开文档
        wps = win32com.client.Dispatch("Kwps.Application")
        wps.Visible = False  # 可选，设置是否显示 WPS 界面
        doc = wps.Documents.Open(os.path.abspath(wps_file_path), ReadOnly=1)
        if wps_file_path == save_path:
            os.remove(wps_file_path)
        # 另存为 docx 文件
        doc.SaveAs(os.path.abspath(save_path), 12)  # 12 表示 docx 格式

        # 关闭文档和 WPS 应用
        doc.Close()
        wps.Quit()
        pythoncom.CoUninitialize()

    @staticmethod
    # 需要安装wps或office
    def doc2docx(input_filepath, output_filepath, keep_active=True):
        pythoncom.CoInitialize()
        input_filepath = Path(input_filepath).resolve()
        output_filepath = Path(output_filepath).resolve()
        word_app = win32com.client.Dispatch("Word.Application")
        doc = word_app.Documents.Open(str(input_filepath))
        try:
            doc.SaveAs2(str(output_filepath), FileFormat=16)
            doc.Close(0)
        except:
            doc.Close(0)

        if not keep_active:
            word_app.Quit()
        pythoncom.CoUninitialize()
    @staticmethod
    # 替换报告单里段落文字:{word}
    # python替换word中的书签变量{word}， 在读取{}时，获取runs不知道为什么会有些连在一起的文字会被拆掉（即使样式一样）
    # 解决：修改{word}里变量名，或在变量名称里加个},比如{word}}
    # obj={}
    # obj["{name}"]=""
    # obj["{age}"]=""
    def replace_txt_in_word(document, obj, logger):
        try:

            # 循环所有段落
            for paragraph in document.paragraphs:
                # run是根据样式来分的，如果文字在一起，但是不同样式，也会分成两个run
                for run in paragraph.runs:
                    # print(run.text)
                    for key, value in obj.items():
                        key = "{}".format(key)
                        if key in run.text:
                            run.text = run.text.replace(key, "{}".format(value))
                            # break
            # 循环所有表格
            for table in document.tables:
                for row in table.rows:
                    for cell in row.cells:
                        cell_text = cell.text
                        # print(cell_text)
                        for key, value in obj.items():
                            key = "{}".format(key)
                            if key in cell_text:
                                for paragraph in cell.paragraphs:
                                    if key in paragraph.text:
                                        for run in paragraph.runs:
                                            run.text = run.text.replace(key, "{}".format(value))


        except Exception as exp:
            logger.error("替换word里的文本失败：" + str(exp))
            # 输出异常信息
            logger.error(exp)
            logger.error(exp.__traceback__.tb_frame.f_globals["__file__"])  # 发生异常所在的文件
            logger.error(exp.__traceback__.tb_lineno)  # 发生异常所在的行数

    @staticmethod
    # 填充报告里的表格
    # obj_list:[{obj}],obj的key与表格的列名一致
    # obj = {}
    # obj["序号"] = ""
    # obj["姓名"] = ""
    # obj["身份证"] = ""
    def fill_table_in_word(table, obj_list, logger, show_max_rows):
        try:
            all_data_list = obj_list
            if len(all_data_list) > 0:

                rows = table.rows
                headCells = rows[0].cells  # 获取头
                for i in range(len(all_data_list)):
                    print("写入" + str(i))
                    # 只输出前***条,避免行数过多，把表格撑爆
                    if i >= show_max_rows:
                        break
                    rowData = all_data_list[i]
                    # 插入表格行数据
                    if i == 0:
                        rowCells_record = table.rows[1].cells
                    else:
                        rowCells_record = table.add_row().cells
                    for index in range(len(headCells)):
                        head = headCells[index]
                        # rowCells_record[index].text = str(rowData.get(head.text))
                        run = rowCells_record[index].paragraphs[0].add_run(str(rowData.get(head.text)))
                        run.font.name = '宋体'
                        run.font.size = Pt(10)
                        r = run._element
                        r.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
                        # 设置表格的对齐方式
                        rowCells_record[index].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                        rowCells_record[index].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

        except Exception as exp:
            logger.error("填充word里的表格失败：" + str(exp))
            # 打印输出异常信息
            logger.error(exp)
            logger.error(exp.__traceback__.tb_frame.f_globals["__file__"])  # 发生异常所在的文件
            logger.error(exp.__traceback__.tb_lineno)  # 发生异常所在的行数

    @staticmethod
    # 替换报告里的的图片
    def replace_image_in_word(document, obj, input_img_path):

        # 循环所有表格
        for table in document.tables:
            for row in table.rows:
                for cell in row.cells:
                    cell_text = cell.text
                    # print(cell_text)
                    for key, value in obj.items():
                        key = "{}".format(key)
                        if key in cell_text:
                            for paragraph in cell.paragraphs:
                                # if key in paragraph.text:
                                #     # 把占位符去掉
                                #     paragraph.text = paragraph.text.replace('{' + key + '}', '')
                                #     run = paragraph.add_run('')
                                #     # 添加图片并指定大小
                                #     if obj[key] != "":
                                #         img_path = CommonHelper.remove_url_head(obj[key])
                                #         input_img = str(settings.BASE_DIR) + "/my_app" + img_path
                                #         print(input_img)
                                #         # img = Image.open(input_img)
                                #         run.add_picture(input_img, height=Cm(1.14))
                                if key in paragraph.text:
                                    has_replaced = False
                                    for run in paragraph.runs:
                                        run.clear()
                                        # 增加转行，使得下面插入图片往下走
                                        run.add_break()
                                        run.add_break()
                                        if not has_replaced:
                                            if obj[key] != "":
                                                print(input_img_path)
                                                if os.path.exists(input_img_path):
                                                    run.add_picture(input_img_path, width=Cm(2.69))
                                                has_replaced = True
